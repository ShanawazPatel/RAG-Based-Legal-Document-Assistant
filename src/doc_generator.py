import io
import logging
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF
from src.utils import get_llm

logging.basicConfig(level=logging.INFO)

DOC_TYPES = [
    "Non-Disclosure Agreement (NDA)",
    "Employment & Service Agreement",
    "Legal Notice",
    "Independent Contractor Agreement",
    "General Affidavit"
]

def generate_legal_draft(doc_type: str, party_a: str, party_b: str, effective_date: str, jurisdiction: str, key_terms: str, provider: str = "Ollama", model_name: str = "llama3.2:1b") -> str:
    """
    Uses AI to generate a complete professional legal document draft.
    """
    prompt = f"""You are an expert Legal Draftsman. Generate a complete, legally binding {doc_type} based on the parameters provided.

DOCUMENT PARAMETERS:
- Document Type: {doc_type}
- Disclosing / First Party (Party A): {party_a}
- Receiving / Second Party (Party B): {party_b}
- Effective Date: {effective_date}
- Governing Jurisdiction / Law: {jurisdiction}
- Key Specific Terms / Scope: {key_terms}

DRAFTING INSTRUCTIONS:
1. Include formal title, opening recitals (WHEREAS), and definitions.
2. Structure into numbered, clear legal clauses (1. Obligations, 2. Confidentiality/Scope, 3. Term & Termination, 4. Governing Law & Dispute Resolution, 5. Severability & Entire Agreement).
3. Include formal signature blocks for Party A and Party B at the end.
4. Ensure standard legal precision and professional formatting.

GENERATE FULL LEGAL DRAFT:"""

    try:
        llm = get_llm(provider=provider, model_name=model_name)
        response = llm.invoke(prompt)
        return response.content.strip() if hasattr(response, "content") else str(response)
    except Exception as e:
        logging.error(f"Error drafting document: {str(e)}")
        return f"Error generating legal document draft: {str(e)}"


def export_to_docx(title: str, text: str) -> io.BytesIO:
    """
    Converts legal draft text into a styled DOCX document buffer.
    """
    doc = DocxDocument()
    
    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = 1 # Center
    
    # Add paragraphs
    for line in text.split("\n"):
        line_clean = line.strip()
        if not line_clean:
            continue
        if line_clean.startswith("# ") or line_clean.startswith("## "):
            doc.add_heading(line_clean.lstrip("#").strip(), level=2)
        else:
            p = doc.add_paragraph(line_clean)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_to_pdf(title: str, text: str) -> bytes:
    """
    Converts legal draft text into a clean PDF using FPDF2.
    """
    class LegalPDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 10)
            self.set_text_color(100, 100, 100)
            self.cell(0, 10, title.upper(), border=False, new_x="LMARGIN", new_y="NEXT", align="C")
            self.line(10, 18, 200, 18)
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}} - Strictly Confidential", border=False, align="C")

    pdf = LegalPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=10)

    for line in text.split("\n"):
        clean_line = line.strip()
        if not clean_line:
            pdf.ln(3)
            continue
        # Encode string safely for standard FPDF latin1
        safe_line = clean_line.encode("latin-1", "replace").decode("latin-1")
        if safe_line.startswith("# ") or safe_line.startswith("## "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 7, safe_line.replace("#", "").strip())
            pdf.set_font("Helvetica", size=10)
        else:
            pdf.multi_cell(0, 6, safe_line)

    return pdf.output()

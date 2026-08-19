import logging
import io
from PyPDF2 import PdfReader
from langchain_ollama import ChatOllama
from fpdf import FPDF
from docx import Document
import streamlit as st

# Logging configuration
logging.basicConfig(level=logging.INFO)

MODEL_NAME = "llama3.2:1b"

def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PdfReader(pdf_file)
        pages = []
        for page in pdf_reader.pages:
            txt = page.extract_text()
            if txt:
                pages.append(txt)
        text = "\n".join(pages)
        return text.strip() if text else None
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {str(e)}")
        return None

def summarize_text(text):
    try:
        llm = ChatOllama(model=MODEL_NAME)
        prompt = f"Summarize the following text in a clear and concise manner:\n\n{text}"
        response = llm.invoke(prompt)
        return response.content.strip() if hasattr(response, "content") else "Error: No summary generated."
    except Exception as e:
        logging.error(f"Error in summarization: {str(e)}")
        return "Error generating summary."

def chat_with_law_bot(user_query):
    try:
        llm = ChatOllama(model=MODEL_NAME)
        response = llm.invoke(user_query)
        return response.content.strip() if hasattr(response, "content") else "Error: No response generated."
    except Exception as e:
        logging.error(f"Error in chatbot: {str(e)}")
        return "Error generating response."

def generate_document(doc_type, user_input):
    if doc_type in ["Legal Notice", "Contract Agreement", "Affidavit"]:
        doc = Document()
        doc.add_heading(doc_type, level=1)
        doc.add_paragraph(user_input)
        return doc
    return None

def download_document(doc, filename, file_type):
    if file_type == "pdf":
        # build plain text from a python-docx Document or fallback to str()
        if hasattr(doc, "paragraphs"):
            text = "\n".join(p.text for p in doc.paragraphs)
        elif isinstance(doc, str):
            text = doc
        else:
            text = str(doc)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        # try to use a unicode-safe font if available, else fall back
        try:
            pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
            pdf.set_font("DejaVu", size=12)
        except Exception:
            pdf.set_font("Arial", size=12)

        for line in text.split("\n"):
            pdf.multi_cell(0, 8, line)

        
        pdf_bytes = pdf.output(dest="S").encode("latin-1")
        st.download_button("📥 Download PDF", data=pdf_bytes, file_name=filename, mime="application/pdf")
    elif file_type == "docx":
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("📥 Download DOCX", data=buffer.getvalue(), file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
